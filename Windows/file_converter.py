import pandas as pd
from docx import Document
import os
import tkinter as tk
from tkinter import filedialog, messagebox
import signal
import sys

# Graceful exit function
def graceful_exit(signal_received=None, frame=None):
    print("\nGracefully shutting down...")
    root.quit()  # Closes the tkinter window
    sys.exit(0)

# Register signal handler for graceful shutdown
signal.signal(signal.SIGINT, graceful_exit)  # Handle Ctrl+C (SIGINT)
signal.signal(signal.SIGTERM, graceful_exit)  # Handle kill command (SIGTERM)

def create_batch_folder():
    # Create a new folder for all converted files
    batch_folder = os.path.join(os.getcwd(), "Converted_Batch_Documents")
    os.makedirs(batch_folder, exist_ok=True)
    return batch_folder

def xlsx_to_csv(xlsx_file, output_dir):
    try:
        # Load the Excel file
        excel_data = pd.ExcelFile(xlsx_file)
        
        # Create a directory for storing CSV files within the batch folder
        file_name = os.path.splitext(os.path.basename(xlsx_file))[0]
        output_subdir = os.path.join(output_dir, f"{file_name}_csv")
        os.makedirs(output_subdir, exist_ok=True)
        
        # Iterate through each sheet in the Excel file
        for sheet_name in excel_data.sheet_names:
            # Read the sheet into a DataFrame
            df = excel_data.parse(sheet_name)
            
            # Generate CSV content
            csv_file = os.path.join(output_subdir, f"{sheet_name}.csv")
            df.to_csv(csv_file, index=False)
            
            print(f"CSV file for sheet '{sheet_name}' created: {csv_file}")
    except Exception as e:
        print(f"Error converting {xlsx_file} to CSV: {e}")

def docx_to_markdown(docx_file, output_dir):
    try:
        # Load the docx file
        doc = Document(docx_file)
        
        # Extract filename without extension and create a markdown file in the batch folder
        file_name = os.path.splitext(os.path.basename(docx_file))[0]
        markdown_file = os.path.join(output_dir, f"{file_name}.md")
        
        # Open a markdown file to write the content with UTF-8 encoding
        with open(markdown_file, "w", encoding="utf-8") as md_file:
            for para in doc.paragraphs:
                # Add a newline after each paragraph for spacing in markdown
                md_file.write(para.text + "\n\n")
            
            # Process tables (if any)
            for table in doc.tables:
                for row in table.rows:
                    row_content = "| " + " | ".join(cell.text for cell in row.cells) + " |\n"
                    md_file.write(row_content)
                # Add a separator between tables for markdown formatting
                md_file.write("\n---\n")
        
        print(f"Markdown file created: {markdown_file}")
    except Exception as e:
        print(f"Error converting {docx_file} to Markdown: {e}")

def convert_files(file_paths):
    try:
        # Create a batch folder for all converted files
        batch_folder = create_batch_folder()
        
        for file_path in file_paths:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == ".xlsx":
                print(f"Converting {file_path} to CSV...")
                xlsx_to_csv(file_path, batch_folder)
            elif file_extension == ".docx":
                print(f"Converting {file_path} to Markdown...")
                docx_to_markdown(file_path, batch_folder)
            else:
                print(f"Unsupported file format for file: {file_path}")
        
        # Show single dialog after all conversions are done
        messagebox.showinfo("Success", "All files have been successfully converted and placed in the 'Converted_Batch_Documents' folder.")
    except Exception as e:
        print(f"Error during batch conversion: {e}")

def select_files():
    # Open file dialog to select multiple files
    file_paths = filedialog.askopenfilenames(filetypes=[("Excel files", "*.xlsx"), ("Word files", "*.docx")])
    if file_paths:
        convert_files(file_paths)

def select_directory():
    # Open directory dialog to select a folder
    directory = filedialog.askdirectory()
    if directory:
        # Find all xlsx and docx files in the directory
        file_paths = []
        for root, dirs, files in os.walk(directory):
            for file in files:
                if file.endswith(".xlsx") or file.endswith(".docx"):
                    file_paths.append(os.path.join(root, file))
        convert_files(file_paths)

# Set up the GUI
root = tk.Tk()
root.title("File Converter")
root.geometry("400x200")

# Add buttons for selecting files or a directory
select_file_button = tk.Button(root, text="Select Files to Convert", command=select_files, padx=10, pady=10)
select_file_button.pack(pady=10)

select_directory_button = tk.Button(root, text="Select Directory to Convert", command=select_directory, padx=10, pady=10)
select_directory_button.pack(pady=10)

# Handle window close event
root.protocol("WM_DELETE_WINDOW", graceful_exit)

# Run the application
try:
    root.mainloop()
except KeyboardInterrupt:
    graceful_exit()
